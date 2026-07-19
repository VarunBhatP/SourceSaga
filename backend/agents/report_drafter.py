"""
Agent: Draft GSOC proposals using a local free explainer.
"""
import os
import uuid
from typing import Dict
from docx import Document
from graph.state import AgentState
from utils.free_explainer import generate_proposal

DOWNLOADS_DIR = "downloads"


def draft_report_agent(state: AgentState) -> Dict:
    """
    Generate formal proposals without any external API.
    """
    print("📝 Agent: Drafting proposals...")
    
    analyses = state.get("analyses", [])
    base_url = os.getenv("API_BASE_URL", "http://localhost:8000")
    
    os.makedirs(DOWNLOADS_DIR, exist_ok=True)
    
    downloads = []
    
    for analysis in analyses:
        context = analysis["context"][:1000]
        plan = analysis["solution_plan"][:800]
        
        print(f"  Drafting proposal for: {analysis['issue_url'][:50]}...")
        proposal_text = generate_proposal(
            context,
            plan,
            analysis["issue_url"],
        )
        
        # Create .docx file
        doc = Document()
        doc.add_heading('Google Summer of Code Project Proposal', level=1)
        
        # Parse and format the proposal
        for line in proposal_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('###'):
                doc.add_heading(line.replace('#', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('**') and line.endswith('**'):
                doc.add_heading(line.replace('**', ''), level=3)
            else:
                doc.add_paragraph(line)
        
        filename = f"proposal_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(DOWNLOADS_DIR, filename)
        doc.save(filepath)
        
        issue_title = analysis["context"].split("\n")[0].replace("**Issue Title:** ", "")
        downloads.append({
            "issue_title": issue_title[:60],
            "download_url": f"{base_url}/api/download/{filename}"
        })
    
    print(f"✅ {len(downloads)} proposals drafted")
    
    return {
        "analyses": analyses,  # ✅ Keep the analyses!
        "report_downloads": downloads,
        "current_step": "reports_ready"
    }

